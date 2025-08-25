from io import BytesIO
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown import markdown
from html2docx import html2docx


# =======================
# CONSTANTES GERAIS
# =======================
DEFAULT_FONT_NAME   = "Calibri"
DEFAULT_FONT_SIZE   = Pt(12)             # corpo do texto
DEFAULT_FONT_COLOR  = RGBColor(0, 0, 0)  # preto

LINE_SPACING        = 1.0                # espaçamento entre linhas (simples)
SPACE_BEFORE        = Pt(0)              # espaço antes do parágrafo
SPACE_AFTER         = Pt(6)              # espaço depois do parágrafo

MARGIN_CM           = 2.5                # todas as margens em cm


def _add_paragraph_border_double(paragraph, color="000000", size_eights_pt=12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'double')
        el.set(qn('w:sz'), str(size_eights_pt))
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '4')
        pBdr.append(el)
    pPr.append(pBdr)


def _add_bottom_border(paragraph, color="444444", size_eights_pt=12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eights_pt))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx_bytes(payload: dict) -> bytes:
    term = (payload.get("term") or "").strip() or "Resultados"
    search_type = (payload.get("search_type") or payload.get("type") or "").capitalize()
    results = payload.get("results") or []

    grouped = defaultdict(list)
    for it in results:
        src = it.get("source") or it.get("book") or it.get("file") or "Geral"
        grouped[src].append(it)

    total_count = sum(len(v) for v in grouped.values())
    doc = Document()

    # ===== Ajustes globais =====
    style = doc.styles["Normal"]
    style.font.name = DEFAULT_FONT_NAME
    style.font.size = DEFAULT_FONT_SIZE
    style.font.color.rgb = DEFAULT_FONT_COLOR

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = SPACE_BEFORE
    pf.space_after = SPACE_AFTER

    for section in doc.sections:
        section.top_margin    = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin   = Cm(MARGIN_CM)
        section.right_margin  = Cm(MARGIN_CM)

    # ===== Cabeçalho =====
    p = doc.add_paragraph(term)
    run = p.runs[0]
    run.font.size = Pt(20)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_paragraph_border_double(p, size_eights_pt=16)

    doc.add_paragraph("")

    # ===== Tipo =====
    p = doc.add_paragraph()
    p.add_run("Tipo de pesquisa: ").bold = True
    p.add_run(search_type)

    doc.add_paragraph("")

    # ===== Termo =====
    p = doc.add_paragraph()
    p.add_run("Termo de pesquisa: ").bold = True
    p.add_run(term)

    doc.add_paragraph("")

    # ===== Estatísticas =====
    p = doc.add_paragraph()
    p.add_run("Total de resultados: ").bold = True
    p.add_run(str(total_count))
    for src, items in grouped.items():
        doc.add_paragraph(f"• {src}: {len(items)}")

    doc.add_paragraph("")

    # ===== Resultados =====
    counter = 1   # <<< INÍCIO DA NUMERAÇÃO GLOBAL

    for src, items in grouped.items():
        doc.add_paragraph("")

        # Badge da fonte (vermelho, bold)
        badge_p = doc.add_paragraph()
        run = badge_p.add_run(f"{src}")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        badge_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Linha divisória
        divider_p = doc.add_paragraph()
        _add_bottom_border(divider_p)

        doc.add_paragraph("")

        for it in items:
            raw_text = it.get("markdown") or it.get("content_text") or it.get("text") or ""
            numbered_md = f"{raw_text}"

            # Markdown -> HTML
            html = markdown(numbered_md, extensions=["extra", "sane_lists"])
            tmp_bytes = html2docx(html, None)
            tmp_doc = Document(tmp_bytes)

            # Copiar parágrafos inteiros com numeração global
            for p in tmp_doc.paragraphs:
                new_p = doc.add_paragraph()

                # Numeração sequencial global
                num_run = new_p.add_run(f"{counter}. ")
                num_run.bold = True
                num_run.font.color.rgb = RGBColor(0, 0, 255)

                # Conteúdo original
                for r in p.runs:
                    new_r = new_p.add_run(r.text)
                    new_r.bold = r.bold
                    new_r.italic = r.italic
                    new_r.font.name = DEFAULT_FONT_NAME
                    new_r.font.size = DEFAULT_FONT_SIZE
                    new_r.font.color.rgb = DEFAULT_FONT_COLOR

                # Justificar o parágrafo
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Garantir formatação de parágrafo
                pf = new_p.paragraph_format
                pf.line_spacing = LINE_SPACING
                pf.space_before = SPACE_BEFORE
                pf.space_after = SPACE_AFTER

                counter += 1

            # Metadados
            meta = [f"Fonte: {src}"]
            if it.get("title"):
                meta.append(f"Título: {it['title']}")
            if it.get("number"):
                meta.append(f"@{it['number']}")
            if it.get("score"):
                try:
                    meta.append(f"Score: {float(it['score']):.3f}")
                except Exception:
                    meta.append(f"Score: {it['score']}")
            if it.get("author"):
                meta.append(f"Autor: {it['author']}")
            if it.get("date"):
                meta.append(f"Data: {it['date']}")
            if it.get("theme"):
                meta.append(f"Tema: {it['theme']}")

            if meta:
                meta_p = doc.add_paragraph(" | ".join(meta))
                meta_p.runs[0].font.size = Pt(9)
                meta_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

                # Garantir formatação de parágrafo
                pf = meta_p.paragraph_format
                pf.line_spacing = LINE_SPACING
                pf.space_before = SPACE_BEFORE
                pf.space_after = SPACE_AFTER

            doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
